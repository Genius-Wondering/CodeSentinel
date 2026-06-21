"""
DocChunker: structure-aware document chunking for PDF, Markdown, TXT, DOCX, and XLSX files.

Strategy:
- Markdown : split on headers (H1/H2/H3), then by character limit within each section.
- PDF      : PDFMiner for layout-aware extraction, then character split with overlap.
- TXT      : RecursiveCharacterTextSplitter directly (no structure to preserve).
- DOCX     : docx2txt extracts paragraph text; character split with overlap.
- XLSX     : openpyxl reads each sheet as a TSV block; each sheet becomes a chunk
             (further split by character if very large).
- Fallback : any other extension → plain RecursiveCharacterTextSplitter.

All chunks get `source_type="doc"` in metadata for retriever filtering.
"""
import os
from typing import List

import openpyxl
from langchain_core.documents import Document
from langchain_text_splitters import (
    MarkdownHeaderTextSplitter,
    RecursiveCharacterTextSplitter,
)

try:
    from langchain_community.document_loaders import PDFMinerLoader
    PDFMINER_AVAILABLE = True
except ImportError:
    from langchain_community.document_loaders import PyPDFLoader as PDFMinerLoader
    PDFMINER_AVAILABLE = False

try:
    from langchain_community.document_loaders import Docx2txtLoader
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

from app.config import config

MARKDOWN_HEADERS = [
    ("#",   "h1"),
    ("##",  "h2"),
    ("###", "h3"),
]

# File extensions supported by this chunker
SUPPORTED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx", ".xlsx", ".xls"}


class DocChunker:
    """
    Splits PDF, Markdown, TXT, DOCX, and XLSX documents into LangChain Documents
    that respect the document's own structure rather than raw character counts.
    """

    def __init__(self):
        self._char_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.DOC_CHUNK_SIZE,
            chunk_overlap=config.DOC_CHUNK_OVERLAP,
        )
        self._md_header_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=MARKDOWN_HEADERS,
            strip_headers=False,
        )

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def chunk(self, file_path: str, doc_id: str, version: int = 1) -> List[Document]:
        """
        Load and chunk a file. Returns a list of Documents with metadata:
            source_type : "doc"
            doc_id      : caller-supplied identifier
            filename    : basename of file_path
            version     : document version number
            page        : page number (PDF) or sheet index (XLSX) or 0
            section     : header breadcrumb (Markdown) or sheet name (XLSX) or ""
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".md":
            docs = self._chunk_markdown(file_path)
        elif ext == ".pdf":
            docs = self._chunk_pdf(file_path)
        elif ext == ".txt":
            docs = self._chunk_txt(file_path)
        elif ext in (".docx",):
            docs = self._chunk_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            docs = self._chunk_xlsx(file_path)
        else:
            docs = self._chunk_generic(file_path)

        base_meta = {
            "source_type": "doc",
            "doc_id": doc_id,
            "filename": os.path.basename(file_path),
            "version": version,
        }
        for doc in docs:
            doc.metadata = {**base_meta, **doc.metadata}

        return docs

    # ------------------------------------------------------------------
    # Markdown
    # ------------------------------------------------------------------

    def _chunk_markdown(self, file_path: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        header_docs = self._md_header_splitter.split_text(text)
        result: List[Document] = []
        for doc in header_docs:
            sub = self._char_splitter.split_documents([doc])
            for chunk in sub:
                breadcrumb = " > ".join(
                    v for k in ("h1", "h2", "h3")
                    if (v := chunk.metadata.get(k))
                )
                chunk.metadata["section"] = breadcrumb
                chunk.metadata["page"] = 0
                result.append(chunk)
        return result

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def _chunk_pdf(self, file_path: str) -> List[Document]:
        loader = PDFMinerLoader(file_path)
        raw_docs = loader.load()

        cleaned: List[Document] = []
        for doc in raw_docs:
            lines = [
                line for line in doc.page_content.split("\n")
                if not _is_boilerplate(line)
            ]
            doc.page_content = "\n".join(lines).strip()
            if doc.page_content:
                cleaned.append(doc)

        chunks = self._char_splitter.split_documents(cleaned)
        for chunk in chunks:
            chunk.metadata["page"] = chunk.metadata.pop(
                "page_number", chunk.metadata.get("page", 0)
            )
            chunk.metadata.setdefault("section", "")
        return chunks

    # ------------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------------

    def _chunk_txt(self, file_path: str) -> List[Document]:
        """Plain text: read with UTF-8 (replace errors), split by character."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        docs = [Document(page_content=text, metadata={"page": 0, "section": ""})]
        return self._char_splitter.split_documents(docs)

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def _chunk_docx(self, file_path: str) -> List[Document]:
        """
        Word document: extract text via docx2txt (preserves paragraph structure),
        then split by character with overlap.
        Falls back to python-docx paragraph join if docx2txt is unavailable.
        """
        if DOCX2TXT_AVAILABLE:
            loader = Docx2txtLoader(file_path)
            raw_docs = loader.load()
        else:
            # python-docx fallback
            import docx as _docx  # type: ignore
            doc = _docx.Document(file_path)
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            raw_docs = [Document(page_content=text, metadata={})]

        chunks = self._char_splitter.split_documents(raw_docs)
        for chunk in chunks:
            chunk.metadata.setdefault("page", 0)
            chunk.metadata.setdefault("section", "")
        return chunks

    # ------------------------------------------------------------------
    # XLSX / XLS
    # ------------------------------------------------------------------

    def _chunk_xlsx(self, file_path: str) -> List[Document]:
        """
        Excel workbook: each sheet is converted to a TSV-style text block.
        Format: "[Sheet: <name>]\n<header row>\n<data rows...>"

        Each sheet becomes at least one Document. Sheets larger than
        DOC_CHUNK_SIZE are further split by the character splitter.
        Empty sheets are skipped.
        """
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        result: List[Document] = []

        for sheet_idx, sheet_name in enumerate(wb.sheetnames):
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if any(c is not None for c in row):
                    rows.append("\t".join(
                        str(c) if c is not None else "" for c in row
                    ))

            if not rows:
                continue  # skip empty sheets

            text = f"[Sheet: {sheet_name}]\n" + "\n".join(rows)
            sheet_doc = Document(
                page_content=text,
                metadata={"page": sheet_idx, "section": sheet_name},
            )
            # Split large sheets further
            sub = self._char_splitter.split_documents([sheet_doc])
            result.extend(sub)

        wb.close()
        return result

    # ------------------------------------------------------------------
    # Generic fallback
    # ------------------------------------------------------------------

    def _chunk_generic(self, file_path: str) -> List[Document]:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        docs = [Document(page_content=text, metadata={"page": 0, "section": ""})]
        return self._char_splitter.split_documents(docs)


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _is_boilerplate(line: str) -> bool:
    """Return True for lines that are noise rather than content."""
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.isdigit():
        return True
    lower = stripped.lower()
    if lower.startswith("page ") and stripped[5:].strip().isdigit():
        return True
    noise_markers = ("www.", "http://", "https://", "©", "copyright", "all rights reserved")
    return any(marker in lower for marker in noise_markers)
